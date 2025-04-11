import docx

# Create a new Word document
doc = docx.Document()

# Add a title
doc.add_heading('Customer Order Information', 0)

# Add customer information
doc.add_paragraph('Customer: Emily Wilson')
doc.add_paragraph('Order ID: ORD24680')
doc.add_paragraph('Date: 2025-04-05')

# Add a table for items
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Quantity'
hdr_cells[2].text = 'Price'

# Add items to the table
items = [
    ('Wireless Headphones', '1', '$129.99'),
    ('Phone Case', '2', '$24.99'),
    ('Screen Protector', '1', '$19.99')
]

for item, qty, price in items:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = qty
    row_cells[2].text = price

# Add total
doc.add_paragraph('Total: $199.96')

# Add shipping information
doc.add_paragraph('Shipping Address:')
doc.add_paragraph('456 Oak Avenue\nSometown, NY 54321')

# Add payment and status
doc.add_paragraph('Payment Method: PayPal')
doc.add_paragraph('Status: Processing')

# Save the document
doc.save('/home/ubuntu/text_extractor/sample_data/sample_order.docx')

print("Sample Word document created successfully.")
