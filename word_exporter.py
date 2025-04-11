"""
Word exporter module.
Exports structured data to Word document format.
"""
import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any


class WordExporter:
    """Exporter for creating Word documents from structured data."""
    
    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the Word exporter with configuration.
        
        Args:
            config: Dictionary containing Word export configurations
        """
        self.config = config.get('word', {})
        self.output_structure = config.get('structure', [])
    
    def export(self, data: List[Dict[str, Any]], output_path: str) -> str:
        """
        Export data to Word document format.
        
        Args:
            data: List of dictionaries containing structured data
            output_path: Directory path where the output file will be saved
            
        Returns:
            Path to the created Word document
        """
        try:
            # Create a new Word document
            doc = docx.Document()
            
            # Add title
            title = self.config.get('title', 'Extracted Data Report')
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add summary if requested
            if self.config.get('include_summary', True):
                doc.add_paragraph(f"Total records: {len(data)}")
                doc.add_paragraph()
            
            # Add table if requested
            if self.config.get('include_table', True) and data:
                # Create mapping of field names to display names
                field_to_display = {
                    item.get('field'): item.get('display_name', item.get('field'))
                    for item in self.output_structure if 'field' in item
                }
                
                # Get ordered list of fields
                ordered_fields = [item.get('field') for item in self.output_structure if 'field' in item]
                
                # If no structure defined, use all fields from first record
                if not ordered_fields and data:
                    ordered_fields = list(data[0].keys())
                
                # Create table
                table = doc.add_table(rows=1, cols=len(ordered_fields))
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                for i, field in enumerate(ordered_fields):
                    display_name = field_to_display.get(field, field)
                    header_cells[i].text = display_name
                    # Make header bold
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for record in data:
                    row_cells = table.add_row().cells
                    for i, field in enumerate(ordered_fields):
                        if field in record:
                            row_cells[i].text = str(record[field])
            
            # Create output file path
            file_name = os.path.join(output_path, 'extracted_data.docx')
            
            # Save the document
            doc.save(file_name)
            
            return file_name
            
        except Exception as e:
            print(f"Error exporting to Word: {str(e)}")
            return ""
